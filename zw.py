from docx import Document
document = Document('/Users/patrickzhang/Documents/yy.docx')
from collections import OrderedDict
import pdb
chinese_lines = []
ipa_lines = []
def detect_chinese(stringA):
    chinese_characters = []
    for character in stringA:
        if character > u'\u4e00' and character < u'\u9fff':
            chinese_characters.append(character)
    return chinese_characters

def detect_ipa(stringA):
    # pdb.set_trace()
    #print(stringA)
    ipa_list = []
    temp_ipa = []
    for num, ipa in enumerate(stringA):
        if ipa == " " or ipa == ".":
            continue
        if not ipa.isdigit():
            temp_ipa.append(ipa)
        else:
            if num<len(stringA)-1:
                if stringA[num+1].isdigit():
                    temp_ipa.append(ipa)
                else:
                    temp_ipa.append(ipa)
                    ipa_list.append(''.join(temp_ipa))
                    temp_ipa = []
            else:
                temp_ipa.append(ipa)
                ipa_list.append(''.join(temp_ipa))
    #print(ipa_list)
    return ipa_list

for num, para in enumerate(document.paragraphs):
    if num%2 == 0:
        chinese_lines.append(para.text)
    else:
        ipa_lines.append(para.text)
chinese_dict = OrderedDict()
for chinese_line,ipa_line in zip(chinese_lines,ipa_lines):
    chinese_list = detect_chinese(chinese_line)
    ipa_list = detect_ipa(ipa_line)
    for chinese_char, ipa_char in zip(chinese_list,ipa_list):
        chinese_dict[chinese_char]=ipa_char
    # print(detect_chinese(chinese_line),detect_ipa(ipa_line))
    # if not len(detect_chinese(chinese_line)) == len(detect_ipa(ipa_line)):
print(chinese_dict)

document = Document()
for key,value in chinese_dict.items():
    document.add_paragraph(
        '{0}\t{1}'.format(key,value)
    )
document.save('./dict.docx')